from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = "/Users/zlzhou/codex/zx-docs/2026年秋算法设计与分析课程教学实施方案.docx"
BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "F2F4F7"
PALE = "E8EEF5"
MUTED = "667085"

def set_font(run, name="PingFang SC", size=11, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for tag, val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn("w:"+tag))
        if node is None: node = OxmlElement("w:"+tag); tcMar.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None: tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i]/1440)
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")
            margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); trPr.append(el)

def cell_text(cell, text, size=9, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.08
    if align is not None: p.alignment = align
    r = p.add_run(text); set_font(r, size=size, bold=bold, color=color)

def add_table(doc, headers, rows, widths, font_size=8.7):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.style = "Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i], LIGHT); cell_text(t.rows[0].cells[i], h, 9, True, DARK, WD_ALIGN_PARAGRAPH.CENTER)
    repeat_header(t.rows[0])
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            if ridx%2: shade(cells[i], "FAFBFC")
            cell_text(cells[i], str(v), font_size, False, None, WD_ALIGN_PARAGRAPH.CENTER if i in (0,1) else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(t,widths)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
    return t

def add_bullet(doc, text, level=0):
    p=doc.add_paragraph(style="List Bullet" if level==0 else "List Bullet 2")
    p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.167
    p.paragraph_format.left_indent=Inches(.5 if level==0 else .75); p.paragraph_format.first_line_indent=Inches(-.25)
    r=p.add_run(text); set_font(r, size=11)
    return p

def add_num(doc, text):
    p=doc.add_paragraph(style="List Number"); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.167
    p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25)
    r=p.add_run(text); set_font(r, size=11); return p

doc=Document()
sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)

styles=doc.styles
normal=styles["Normal"]; normal.font.name="PingFang SC"; normal.font.size=Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"),"PingFang SC")
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for key,size,color,before,after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK,8,4)]:
    s=styles[key]; s.font.name="PingFang SC"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s._element.rPr.rFonts.set(qn("w:eastAsia"),"PingFang SC")
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

# Running furniture
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
hr=hp.add_run("2026年秋 · 算法设计与分析"); set_font(hr,size=9,color=MUTED)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=fp.add_run("课程教学实施方案"); set_font(fr,size=9,color=MUTED)
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); fp._p.append(fld)

# Workshop-agenda opening block
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
r=p.add_run("COURSE IMPLEMENTATION PLAN"); set_font(r,size=10,bold=True,color=BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
r=p.add_run("2026年秋算法设计与分析课程教学实施方案"); set_font(r,size=25,bold=True,color=DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18)
r=p.add_run("以 LeetCode 热题100为主线，以蓝桥杯、洛谷典型题为补充的实战型教学设计"); set_font(r,size=12,color=MUTED)

meta=[("总学时","48学时"),("建议组织","16次课 × 3学时"),("课堂主线","现场分析 · 编码 · 测试 · 复盘"),("课外训练","每日不少于2题")]
t=doc.add_table(rows=1,cols=4); t.style="Table Grid"
for i,(a,b) in enumerate(meta):
    shade(t.rows[0].cells[i],PALE); c=t.rows[0].cells[i]; c.text=""
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
    set_font(p.add_run(a),size=8.5,bold=True,color=BLUE)
    p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
    set_font(p.add_run(b),size=10,bold=True,color=DARK)
set_table_geometry(t,[2340]*4)

doc.add_heading("一、课程定位与设计思路",level=1)
p=doc.add_paragraph("本课程面向已具备一门程序设计语言基础的学生，目标不是停留在算法概念的记忆与证明，而是通过高频、可迁移的实战训练，使学生形成“识别问题结构—选择算法范式—论证正确性—分析复杂度—实现与调试—复盘迁移”的完整能力链。课程以 LeetCode 热题100的知识结构为主线，精选蓝桥杯、洛谷等平台中难度和背景适宜的题目进行迁移训练。")
p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
p=doc.add_paragraph("核心原则：每节课尽可能多地接触典型题，但不追求简单刷题数量；每道核心题必须沉淀为可复用的思维模板、代码模板和错误清单。48学时结束时，学生应能独立完成常见中等难度算法题，并能清楚说明算法选择、正确性依据与时间/空间复杂度。")
p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading("二、课程教学目标",level=1)
objectives=[
"知识目标：掌握哈希、双指针、滑动窗口、子串、数组、矩阵、链表、二叉树、图论、回溯、二分查找、栈、堆、贪心、动态规划和常用技巧等核心知识。",
"能力目标：能够将自然语言问题抽象为数据结构与状态关系，选择合适算法，独立编写可运行代码，并用测试用例定位边界错误。",
"分析目标：能够解释算法正确性，计算时间复杂度与空间复杂度，比较不同解法的适用条件与工程权衡。",
"迁移目标：能够把课堂形成的模式迁移到 LeetCode、蓝桥杯、洛谷及课程综合任务中，解决未见过但结构相似的问题。",
"学习习惯目标：形成每日稳定训练、错题归档、代码复盘和阶段总结的习惯，逐步建立个人算法模板库。"
]
for x in objectives:add_bullet(doc,x)

doc.add_heading("三、教学方法与课堂组织",level=1)
doc.add_heading("（一）问题驱动的现场解题",level=2)
p=doc.add_paragraph("教师以真实平台题目作为知识入口，先隐藏标准答案，由学生识别输入输出、约束条件、数据规模和边界，再现场完成从朴素方案到优化方案的推导。板书或投屏必须保留关键思路演化过程，使学生看到算法是如何“长出来”的。")
doc.add_heading("（二）一课一闭环：讲、练、评、结",level=2)
steps=[
"问题热身（10—15分钟）：用1道易题或旧题变式激活先修知识。",
"方法精讲（30—40分钟）：用1道代表题讲清数据结构、算法范式、正确性和复杂度。",
"现场编码（35—45分钟）：教师边写边解释变量含义、不变量、边界处理和调试方法。",
"同类迁移（45—55分钟）：学生独立或结对完成1—2道同构/变式题，教师巡回答疑。",
"挑战提升（20—30分钟）：选讲竞赛题或中等难度题，比较多种解法。",
"复盘小结（10—15分钟）：形成“识别信号—解题模板—易错点—复杂度”四项总结，并布置每日题。"
]
for s in steps:add_num(doc,s)
doc.add_heading("（三）多平台题源的分工",level=2)
rows=[("LeetCode 热题100","课程知识主线与典型面试题","结构稳定、模式清晰，便于形成算法范式"),("蓝桥杯","综合应用、模拟与竞赛限时训练","强化读题、实现速度和综合建模"),("洛谷","专题训练与难度分层","题量丰富，适合从模板题到提高题的梯度练习")]
add_table(doc,["题源","主要用途","使用原则"],rows,[1800,2700,4860],9.2)

doc.add_heading("四、教学内容与48学时进度安排",level=1)
p=doc.add_paragraph("建议按16次课组织，每次3学时。下表以热题100的模块顺序为主，并根据教学逻辑合并相近主题。题目可依据学生语言基础和平台更新进行等价替换，但每次课的核心模式与能力目标应保持稳定。")
schedule=[
(1,"导论、复杂度与解题流程","两数之和、最长连续序列；补充：枚举优化","建立六步解题法；掌握哈希空间换时间","2讲解+1练习"),
(2,"双指针","移动零、盛最多水的容器、三数之和","掌握相向/同向指针与去重","2讲解+2练习"),
(3,"滑动窗口与子串","无重复字符的最长子串、找到字符串中所有字母异位词","理解窗口不变量、扩张与收缩条件","2讲解+2练习"),
(4,"数组与区间技巧","最大子数组、合并区间、轮转数组、除自身以外数组的乘积","掌握前缀/后缀、区间排序与状态维护","3讲解+1练习"),
(5,"矩阵与模拟","矩阵置零、螺旋矩阵、旋转图像、搜索二维矩阵","提升下标控制、分层遍历与原地修改能力","3讲解+1练习"),
(6,"链表基础与综合","相交链表、反转链表、回文链表、环形链表、合并有序链表","掌握虚拟头结点、快慢指针、指针重连","3讲解+2练习"),
(7,"链表进阶与栈","两两交换、K个一组翻转、随机链表复制；有效括号","处理复杂指针关系，理解栈的配对模型","3讲解+1练习"),
(8,"二叉树遍历与结构","中序遍历、最大深度、翻转二叉树、对称二叉树、直径","掌握递归定义、返回值设计和遍历框架","3讲解+2练习"),
(9,"二叉树进阶","层序遍历、验证搜索树、第K小元素、最近公共祖先、路径总和","掌握BFS、BST性质与树上信息汇总","3讲解+1练习"),
(10,"图论基础","岛屿数量、腐烂的橘子、课程表、实现Trie","掌握网格DFS/BFS、拓扑排序与前缀树","3讲解+1练习"),
(11,"回溯","全排列、子集、组合总和、括号生成、单词搜索","掌握选择—递归—撤销框架与剪枝","3讲解+2练习"),
(12,"二分查找与搜索空间","二分查找、搜索旋转排序数组、寻找峰值、搜索二维矩阵II","能定义单调性、区间和循环不变量","3讲解+1练习"),
(13,"栈、单调栈与堆","最小栈、每日温度、柱状图最大矩形、数组第K大、前K高频元素","掌握单调结构和Top-K模型","3讲解+2练习"),
(14,"贪心与综合技巧","买卖股票最佳时机、跳跃游戏、划分字母区间；只出现一次的数字","识别局部最优与全局最优，理解位运算","3讲解+1练习"),
(15,"动态规划基础","爬楼梯、打家劫舍、完全平方数、零钱兑换、单词拆分","掌握状态、转移、初始化、遍历顺序","3讲解+2练习"),
(16,"动态规划进阶与综合考核","最长递增子序列、最长公共子序列、编辑距离；限时综合题","形成一维/二维DP框架，完成综合迁移与课程复盘","2讲解+综合测评")]
add_table(doc,["次","主题","代表性实战题（示例）","关键能力","课堂产出"],schedule,[520,1300,3370,2820,1350],7.8)

doc.add_heading("五、LeetCode 热题100内容模块建议",level=1)
mods=[("基础数据处理","哈希、双指针、滑动窗口、子串、数组、矩阵","从数据规模判断复杂度；维护集合、窗口、区间与下标不变量"),("线性结构","链表、栈、单调栈、堆","指针操作、后进先出、候选集维护与Top-K"),("非线性结构","二叉树、图、Trie","递归、遍历、连通性、拓扑关系与层次结构"),("搜索与生成","二分查找、回溯","利用单调性缩小空间；系统枚举并剪枝"),("优化方法","贪心、动态规划","建立局部选择依据；定义状态和状态转移"),("综合技巧","位运算、排序、前缀和、边界与工程实现","提高代码简洁性、鲁棒性与复杂度意识")]
add_table(doc,["能力板块","主要内容","教学关注点"],mods,[1600,3000,4760],9)

doc.add_heading("六、学生练习与学习管理",level=1)
doc.add_heading("（一）每日两题制度",level=2)
for x in [
"基础题1道：与当周课堂模板高度同构，要求独立完成并通过全部测试。",
"迁移题1道：改变数据表达、约束或问题背景，要求写出思路、复杂度与至少1个自拟边界用例。",
"建议在16次课期间设置不少于32道必做题；如按完整教学周持续执行，可扩展为“2题/学习日”，由教师按周发布清单并控制总负荷。",
"每位学生维护个人错题本，至少记录：错误类型、失败用例、修正后的关键不变量、可迁移模板。"
]:add_bullet(doc,x)
doc.add_heading("（二）分层练习",level=2)
rows=[("A层：保底","平台简单题/核心模板题","所有学生必须完成；关注正确实现与复杂度达标"),("B层：标准","热题100中等题/常见变式","课程主体；要求独立分析并限时完成"),("C层：挑战","蓝桥杯、洛谷提高题或综合题","供学有余力者；鼓励多解法、优化与讲题")]
add_table(doc,["层级","题目类型","要求"],rows,[1600,3000,4760],9.2)
doc.add_heading("（三）提交与反馈",level=2)
for x in ["提交内容包括可运行代码、核心思路、复杂度、自测用例；不接受只有通过截图而无解释的提交。","教师每周抽取共性错误进行10分钟集中复盘；优秀解法由学生在课堂讲解，形成同伴教学。","对连续未完成、照搬代码或基础模板掌握不牢的学生，安排短时面谈和针对性补练。"]:add_bullet(doc,x)

doc.add_heading("七、考核与评价建议",level=1)
assess=[("平时每日题与作业","30%","完成度、正确性、复杂度说明、错题订正；兼顾持续性"),("课堂参与与随堂练习","15%","现场分析、编码、测试、讨论与讲题表现"),("阶段测验","20%","两次限时上机，覆盖基础模板和同类迁移"),("课程综合项目/大作业","15%","选择一个综合问题，提交设计、代码、测试与复盘报告"),("期末上机考核","20%","2—3题，考查独立建模、实现、复杂度与鲁棒性")]
add_table(doc,["评价项目","建议权重","评价要点"],assess,[2300,1200,5860],9)
p=doc.add_paragraph("评价原则：过程性评价与结果性评价结合；代码是否通过只是基础，还要评价解释能力、复杂度意识、边界测试和迁移能力。对学术诚信应明确要求，允许讨论思路，但提交代码必须能够独立解释。")

doc.add_heading("八、课程综合任务建议",level=1)
p=doc.add_paragraph("综合任务可采用“专题算法包”或“限时题解集”形式。学生从一个主题中选择3—5道递进题，提交问题抽象、朴素解法、优化过程、正确性说明、复杂度、代码、测试数据和复盘。鼓励将同一算法用于不同平台题目，以证明迁移能力。")
for x in ["选题示例1：滑动窗口专题——定长窗口、变长窗口、计数窗口与最优区间。","选题示例2：树与图遍历专题——递归、BFS、连通块、拓扑排序。","选题示例3：动态规划专题——线性DP、背包、序列DP与状态压缩入门。","展示要求：8分钟讲解+现场问答；必须解释至少一个错误方案或性能瓶颈。"]:add_bullet(doc,x)

doc.add_heading("九、教学质量保障与实施条件",level=1)
checks=[
"课前：教师完成题目难度、前置知识、边界用例与多语言代码验证；准备主解法和备选解法。",
"课中：控制讲授与练习比例，原则上学生动手时间不少于每次课的三分之一；实时收集通过率和错误类型。",
"课后：依据提交数据调整下一次课的热身题和补救内容；对高频错误形成班级错误清单。",
"环境：统一编程语言版本、在线判题账号和代码模板；课堂配备可投屏IDE与稳定网络，并准备离线题面和测试数据。",
"阶段复盘：第4、8、12、16次课进行能力盘点，比较正确率、平均用时、独立完成率和复杂度说明质量。"
]
for x in checks:add_bullet(doc,x)

doc.add_heading("十、预期学习成果",level=1)
p=doc.add_paragraph("完成48学时及配套训练后，学生应能够：面对常见算法问题，在合理时间内识别主要模式；独立实现并调试典型算法；对解法进行正确性和复杂度说明；完成从课堂例题到竞赛/平台变式题的迁移；形成可持续的算法训练方法。建议以“中等题独立完成率、首次通过率、平均解题时间、错题复发率、口头解释质量”作为课程成效的核心观测指标。")

doc.add_heading("附录：单题复盘模板",level=1)
review=[("题目与来源","题目名称、平台、编号或链接"),("识别信号","数据规模、关键词、结构特征"),("核心思路","为何选择该数据结构/算法范式"),("关键不变量","循环、窗口、递归或DP过程中始终成立的条件"),("复杂度","时间复杂度、空间复杂度及主要来源"),("失败记录","错误用例、错误原因、修复方式"),("迁移方向","可改变哪些条件形成新题；对应哪些同类题")]
add_table(doc,["项目","记录要求"],review,[1900,7460],9.4)

doc.core_properties.title="2026年秋算法设计与分析课程教学实施方案"
doc.core_properties.subject="算法设计与分析课程教学方法、教学内容与48学时安排"
doc.core_properties.author="课程组"
doc.save(OUT)
print(OUT)
