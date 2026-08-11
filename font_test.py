from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
d=Document()
p=d.add_paragraph("默认字体：算法设计与分析课程教学实施方案")
for f in ["Songti SC","Heiti SC","Hiragino Sans GB","PingFang SC","Arial Unicode MS","STSong","SimSun","Noto Sans CJK SC"]:
 p=d.add_paragraph(); r=p.add_run(f+"：算法设计与分析课程教学实施方案")
 r.font.name=f; r.font.size=Pt(18)
 for k in ["ascii","hAnsi","eastAsia","cs"]: r._element.get_or_add_rPr().rFonts.set(qn("w:"+k),f)
d.save("/Users/zlzhou/codex/zx-docs/font_test.docx")
